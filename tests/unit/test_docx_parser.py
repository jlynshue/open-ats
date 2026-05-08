"""Unit tests for the DOCX resume parser (FR-1.2).

Each test maps to one acceptance criterion in the Sprint 2 contract
(``.context/sprints/02-contract.md``) or PRD §FR-1.2.

Test fixtures (.docx) are generated from the markdown originals by
``tests/fixtures/_build_binary_fixtures.py``. They are checked into
git so this test module doesn't need to regenerate them.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from open_ats.models import Resume, SectionType
from open_ats.parsers import (
    DocxParser,
    ResumeParseError,
    UnsupportedFormatError,
    parse_resume,
)
from open_ats.parsers.docx_parser import default_docx_parser

EXPECTED_CONTACTS = {
    "entry_level.docx": {
        "full_name": "Avery Chen",
        "email": "avery.chen.entry@example.com",
    },
    "mid_level.docx": {
        "full_name": "Jordan Patel",
        "email": "jordan.patel.mid@example.com",
    },
    "executive.docx": {
        "full_name": "Morgan Reyes",
        "email": "morgan.reyes.exec@example.com",
    },
}

ALL_FIXTURES = tuple(EXPECTED_CONTACTS.keys())


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


# ─── Threshold #2/#3 — DOCX dispatch ─────────────────────────────────


@pytest.mark.parametrize("fixture", ALL_FIXTURES)
def test_dispatcher_routes_docx(fixture: str, resumes_dir: Path) -> None:
    """parse_resume() routes .docx through DocxParser and tags source_format."""
    resume = parse_resume(resumes_dir / fixture)
    assert isinstance(resume, Resume)
    assert resume.source_format == "docx"


# ─── Threshold #4 — DOCX coverage ────────────────────────────────────


@pytest.mark.parametrize("fixture", ALL_FIXTURES)
def test_contact_extracted(fixture: str, resumes_dir: Path, parser: DocxParser) -> None:
    """Every fixture yields a Contact with full_name + email both non-None."""
    resume = parser.parse(resumes_dir / fixture)
    assert resume.contact.full_name == EXPECTED_CONTACTS[fixture]["full_name"]
    assert resume.contact.email == EXPECTED_CONTACTS[fixture]["email"]


@pytest.mark.parametrize("fixture", ALL_FIXTURES)
def test_section_detection_at_least_three(
    fixture: str, resumes_dir: Path, parser: DocxParser
) -> None:
    """≥3 of {SUMMARY, EXPERIENCE, EDUCATION, SKILLS, PROJECTS} per fixture."""
    resume = parser.parse(resumes_dir / fixture)
    targets = {
        SectionType.SUMMARY,
        SectionType.EXPERIENCE,
        SectionType.EDUCATION,
        SectionType.SKILLS,
        SectionType.PROJECTS,
    }
    detected = {s.type for s in resume.sections} & targets
    assert len(detected) >= 3, f"{fixture}: detected only {detected}"


def test_experience_entries_split_via_bold_runs(
    resumes_dir: Path, parser: DocxParser
) -> None:
    """Bold runs should be preserved as **markers** so MarkdownParser
    splits experience entries correctly."""
    resume = parser.parse(resumes_dir / "mid_level.docx")
    # mid_level has 3 distinct roles; bold-preservation must surface them all.
    assert len(resume.experience) == 3
    for entry in resume.experience:
        assert entry.title and entry.title != "(unknown)"
        assert entry.company and entry.company != "(unknown)"


# ─── Threshold #6 — Table warning ────────────────────────────────────


def test_table_emits_warning(tmp_path: Path, parser: DocxParser) -> None:
    """A DOCX containing a table surfaces formatting.table_detected."""
    doc_path = tmp_path / "with_table.docx"
    document = docx.Document()
    document.add_heading("Test Table User", level=1)
    document.add_paragraph("test.table@example.com")
    document.add_heading("Summary", level=2)
    document.add_paragraph("A short summary.")
    document.add_heading("Experience", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Engineer"
    table.cell(0, 1).text = "Acme Co"
    table.cell(1, 0).text = "2020 – 2023"
    table.cell(1, 1).text = "Built things"
    document.save(str(doc_path))

    resume = parser.parse(doc_path)
    codes = [w.code for w in resume.parser_warnings]
    assert "formatting.table_detected" in codes


def test_table_cells_not_lost(tmp_path: Path, parser: DocxParser) -> None:
    """Even when a table is flagged, its cell text reaches the rendered body
    so contact info / dates inside the table aren't silently dropped."""
    doc_path = tmp_path / "table_with_email.docx"
    document = docx.Document()
    document.add_heading("Cell Email Person", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Email:"
    table.cell(0, 1).text = "cell.email@example.com"
    document.add_heading("Experience", level=2)
    document.add_paragraph("Something.")
    document.save(str(doc_path))

    resume = parser.parse(doc_path)
    assert resume.contact.email == "cell.email@example.com"


# ─── Threshold #11 — Determinism ─────────────────────────────────────


@pytest.mark.parametrize("fixture", ALL_FIXTURES)
def test_parse_is_deterministic(
    fixture: str, resumes_dir: Path, parser: DocxParser
) -> None:
    a = parser.parse(resumes_dir / fixture).model_dump_json()
    b = parser.parse(resumes_dir / fixture).model_dump_json()
    assert a == b


# ─── Threshold #12 — Format dispatch / unsupported PDF still rejected ─


def test_pdf_still_unsupported(tmp_path: Path) -> None:
    fake = tmp_path / "resume.pdf"
    fake.write_bytes(b"%PDF-1.4 fake")
    with pytest.raises(UnsupportedFormatError):
        parse_resume(fake)


# ─── Error handling ──────────────────────────────────────────────────


def test_corrupt_docx_raises_resume_parse_error(
    tmp_path: Path, parser: DocxParser
) -> None:
    """A non-DOCX bytestream with a .docx extension raises ResumeParseError."""
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not really a docx package")
    with pytest.raises(ResumeParseError):
        parser.parse(bad)


def test_empty_docx_raises_resume_parse_error(
    tmp_path: Path, parser: DocxParser
) -> None:
    """A valid but contentless DOCX raises ResumeParseError."""
    doc_path = tmp_path / "empty.docx"
    docx.Document().save(str(doc_path))
    with pytest.raises(ResumeParseError):
        parser.parse(doc_path)


# ─── Default singleton ───────────────────────────────────────────────


def test_default_singleton_stable() -> None:
    from open_ats.parsers.docx_parser import default_docx_parser as again

    assert default_docx_parser is again
