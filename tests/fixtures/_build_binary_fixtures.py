"""Regenerate DOCX, TXT, and PDF resume fixtures from the markdown originals.

Run once per fixture refresh:

    python tests/fixtures/_build_binary_fixtures.py

Reads ``tests/fixtures/resumes/{entry_level,mid_level,executive}.md`` and
writes the same content as ``.docx`` (python-docx) and ``.txt``
(plain-text demarkdown). Additionally builds three PDF fixtures via
fpdf2:

- ``single_column.pdf`` — entry_level content, single-column layout.
- ``multi_column.pdf`` — same content, two-column layout (degraded
  extraction expected; the parser must still recover ≥2 sections).
- ``image_based.pdf`` — single page with no text content (just a
  drawn rectangle), used to verify the FR-1.3 image-based-PDF error
  path. pdfplumber returns empty text from this fixture.

The generated files are checked into git so tests are deterministic
without running this script in CI.

This is a *fixtures* utility, not part of the runtime package — it
imports python-docx + fpdf2 and may be skipped in environments where
the dev extras are missing.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.document import Document as DocumentT
from fpdf import FPDF

FIXTURES_DIR = Path(__file__).parent / "resumes"
SOURCES = ("entry_level", "mid_level", "executive")

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
_BOLD_TITLE_RE = re.compile(r"^\*\*(.+?)\*\*\s*(.*)$")


def main() -> None:
    for stem in SOURCES:
        md_path = FIXTURES_DIR / f"{stem}.md"
        if not md_path.exists():
            raise FileNotFoundError(md_path)
        markdown = md_path.read_text(encoding="utf-8")
        _write_docx(markdown, FIXTURES_DIR / f"{stem}.docx")
        _write_txt(markdown, FIXTURES_DIR / f"{stem}.txt")
        print(f"  generated {stem}.docx and {stem}.txt")

    # PDF fixtures use the entry_level markdown body; multi-column
    # exercises pdfplumber's column-handling, image-based exercises
    # the empty-extraction error path.
    entry_md = (FIXTURES_DIR / "entry_level.md").read_text(encoding="utf-8")
    _write_pdf_single_column(entry_md, FIXTURES_DIR / "single_column.pdf")
    _write_pdf_multi_column(entry_md, FIXTURES_DIR / "multi_column.pdf")
    _write_pdf_image_based(FIXTURES_DIR / "image_based.pdf")
    print("  generated single_column.pdf, multi_column.pdf, image_based.pdf")


# ─── DOCX ─────────────────────────────────────────────────────────────


def _write_docx(markdown: str, out: Path) -> None:
    document: DocumentT = docx.Document()
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        match = _HEADING_RE.match(line)
        if match:
            depth = len(match.group(1))
            document.add_heading(
                _strip_inline_markdown(match.group(2)), level=min(depth, 9)
            )
            continue
        bullet = _BULLET_RE.match(line)
        if bullet:
            document.add_paragraph(
                _strip_inline_markdown(bullet.group(1)), style="List Bullet"
            )
            continue
        bold_title = _BOLD_TITLE_RE.match(line.lstrip())
        if bold_title:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_strip_inline_markdown(bold_title.group(1)))
            run.bold = True
            tail = bold_title.group(2)
            if tail:
                paragraph.add_run(" " + _strip_inline_markdown(tail))
            continue
        document.add_paragraph(_strip_inline_markdown(line))
    document.save(str(out))


# ─── TXT ──────────────────────────────────────────────────────────────


def _write_txt(markdown: str, out: Path) -> None:
    """Render markdown as plain text suitable for the TXT parser.

    Conversion rules:
      - H1 ``# Name`` → name on its own line (no marker)
      - H2/H3 ``## Foo`` → ``FOO`` (ALL CAPS) on its own line
      - bullets ``- Foo`` → ``- Foo`` (preserved; the TXT parser
        recognises ``-``)
      - **bold** runs → plain text
      - other lines pass through
    """
    out_lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        match = _HEADING_RE.match(line)
        if match:
            depth = len(match.group(1))
            heading = _strip_inline_markdown(match.group(2))
            if depth == 1:
                out_lines.append(heading)
            else:
                out_lines.append(heading.upper())
            continue
        bullet = _BULLET_RE.match(line)
        if bullet:
            out_lines.append(f"- {_strip_inline_markdown(bullet.group(1))}")
            continue
        out_lines.append(_strip_inline_markdown(line))
    out.write_text("\n".join(out_lines) + "\n", encoding="utf-8")


# ─── PDF ──────────────────────────────────────────────────────────────


# Page-layout constants (millimetres). fpdf2 defaults to A4 portrait.
_PAGE_W = 210.0
_MARGIN = 15.0
_COLUMN_GAP = 10.0
_FONT_PT = 11
_LINE_H = 6.0


def _new_pdf() -> FPDF:
    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=_MARGIN)
    pdf.add_page()
    pdf.set_font("Helvetica", size=_FONT_PT)
    pdf.set_margins(_MARGIN, _MARGIN, _MARGIN)
    return pdf


def _write_pdf_single_column(markdown: str, out: Path) -> None:
    """Single-column resume PDF with markdown headings flattened."""
    pdf = _new_pdf()
    usable_width = _PAGE_W - 2 * _MARGIN
    for raw_line in markdown.splitlines():
        _emit_pdf_line(pdf, raw_line, width=usable_width)
    pdf.output(str(out))


def _write_pdf_multi_column(markdown: str, out: Path) -> None:
    """Two-column layout. pdfplumber will interleave columns by line; the
    parser must still recover ≥2 sections (degraded but functional).
    """
    pdf = _new_pdf()
    column_width = (_PAGE_W - 2 * _MARGIN - _COLUMN_GAP) / 2

    # Split the markdown roughly in half by lines so each column gets
    # contiguous content. The contact header always lives in the left
    # column.
    lines = markdown.splitlines()
    midpoint = len(lines) // 2
    left, right = lines[:midpoint], lines[midpoint:]

    top_y = pdf.get_y()
    # Left column.
    pdf.set_xy(_MARGIN, top_y)
    for raw_line in left:
        pdf.set_x(_MARGIN)
        _emit_pdf_line(pdf, raw_line, width=column_width)
    bottom_left = pdf.get_y()

    # Right column starts at the same y as the left column.
    right_x = _MARGIN + column_width + _COLUMN_GAP
    pdf.set_xy(right_x, top_y)
    for raw_line in right:
        pdf.set_x(right_x)
        _emit_pdf_line(pdf, raw_line, width=column_width)
    bottom_right = pdf.get_y()

    # Park the cursor below both columns.
    pdf.set_y(max(bottom_left, bottom_right))
    pdf.output(str(out))


def _write_pdf_image_based(out: Path) -> None:
    """A single page with no text — only a drawn rectangle.

    pdfplumber's ``extract_text()`` returns an empty string for this
    fixture, which is the signal :class:`PdfParser` uses to raise the
    FR-1.3 ``image-based`` error. Functionally equivalent to a scanned
    document with no OCR layer, without needing Pillow as a dev dep.
    """
    pdf = FPDF(unit="mm", format="A4")
    pdf.add_page()
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.4)
    pdf.rect(_MARGIN, _MARGIN, _PAGE_W - 2 * _MARGIN, 60.0)
    pdf.output(str(out))


_UNICODE_REPLACEMENTS = {
    "—": "-",  # em-dash → hyphen
    "–": "-",  # en-dash → hyphen
    "•": "*",  # bullet → asterisk
    "“": '"',  # curly double quote (left)
    "”": '"',  # curly double quote (right)
    "‘": "'",  # curly single quote (left)
    "’": "'",  # curly single quote (right)
    "…": "...",  # ellipsis
}


def _to_helvetica_safe(text: str) -> str:
    """Substitute Unicode characters not supported by Helvetica's WinAnsi
    encoding. fpdf2 ships only the 14 Adobe core fonts by default; using
    a TrueType font would require shipping it. The substitutions below
    keep readability without adding a font dep."""
    for src, dst in _UNICODE_REPLACEMENTS.items():
        if src in text:
            text = text.replace(src, dst)
    return text


def _emit_pdf_line(pdf: FPDF, raw_line: str, *, width: float) -> None:
    """Render a single line of markdown into a PDF, preserving
    enough structural cues for pdfplumber → text → heuristic to
    re-recognise headings + bullets."""
    line = _to_helvetica_safe(raw_line.rstrip())

    if not line:
        # Paragraph break; advance one line height.
        pdf.ln(_LINE_H)
        return

    heading = _HEADING_RE.match(line)
    if heading:
        depth = len(heading.group(1))
        text = _strip_inline_markdown(heading.group(2))
        # Render H1 as the candidate's name (default font, mixed case).
        # Render H2+ as ALL-CAPS bold so pdfplumber → text yields a
        # heading line our heuristic recognises.
        if depth == 1:
            pdf.set_font("Helvetica", "B", _FONT_PT + 4)
            pdf.multi_cell(width, _LINE_H, text)
            pdf.set_font("Helvetica", size=_FONT_PT)
        else:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", _FONT_PT)
            pdf.multi_cell(width, _LINE_H, text.upper())
            pdf.set_font("Helvetica", size=_FONT_PT)
            pdf.ln(1)
        return

    bullet = _BULLET_RE.match(line)
    if bullet:
        text = _strip_inline_markdown(bullet.group(1))
        pdf.multi_cell(width, _LINE_H, f"- {text}")
        return

    text = _strip_inline_markdown(line)
    pdf.multi_cell(width, _LINE_H, text)


def _strip_inline_markdown(text: str) -> str:
    """Drop **bold** and `code` decorations from a single line."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


if __name__ == "__main__":
    main()
